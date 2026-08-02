"""Templates WhatsApp para aprovação Meta — com botões Quick Reply ao máximo."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path(r"C:\site-record\docs\WhatsApp_Templates_Aprovacao_Meta.docx")


def _font(run, size: int = 11, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _p(doc: Document, text: str = "", *, bold: bool = False, size: int = 11, space_after: int = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    _font(p.add_run(text), size=size, bold=bold)


def _h(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        _font(run, size=16 if level == 1 else 13, bold=True)


def _code(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(line if line else "")
        _font(run, size=11)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        if not line:
            p.add_run("\u00a0")


def _kv(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _font(p.add_run(f"{label}: "), bold=True)
    _font(p.add_run(value))


def _template(
    doc: Document,
    *,
    titulo: str,
    nome: str,
    finalidade: str,
    body: list[str],
    botoes: list[str],
    variaveis: list[str],
    notas: list[str] | None = None,
    tipo_botoes: str = "QUICK_REPLY",
) -> None:
    _h(doc, titulo, 2)
    _kv(doc, "Nome sugerido", nome)
    _kv(doc, "Categoria", "UTILITY")
    _kv(doc, "Idioma", "pt_BR")
    _kv(doc, "Finalidade", finalidade)
    _p(doc, "Corpo (colar no Manager — cada linha = Enter):", bold=True)
    _code(doc, body)
    doc.add_paragraph()
    _p(doc, f"Botões — obrigatório ({tipo_botoes}):", bold=True)
    for i, b in enumerate(botoes, start=1):
        _p(doc, f"  {i}. {b}  (≤25 caracteres)", space_after=2)
    _p(doc, "Variáveis (CRM):", bold=True)
    for v in variaveis:
        _p(doc, f"• {v}", space_after=2)
    if notas:
        _p(doc, "Notas:", bold=True)
        for n in notas:
            _p(doc, f"• {n}", space_after=2)
    doc.add_paragraph()


# Abertura padrão: {{1}} saudação · {{2}} primeiro nome

T1 = [
    "Olá, {{1}}!",
    "",
    "{{2}},",
    "",
    "📋 *RESUMO DO PEDIDO NIO FIBRA*",
    "",
    "👤 Cliente: {{3}}",
    "CPF: {{4}}",
    "E-mail: {{5}}",
    "",
    "📍 Endereço:",
    "CEP: {{6}}",
    "Logradouro: {{7}}",
    "Número: {{8}}",
    "Complemento: {{9}}",
    "Bairro: {{10}}",
    "Cidade: {{11}}",
    "",
    "💳 Pagamento: {{12}}",
    "📦 Plano: {{13}}",
    "📅 Fidelidade: {{14}}",
    "",
    "💰 Taxa de habilitação:",
    "Conforme o contrato, a taxa de habilitação fica isenta",
    "quando cumprida a fidelidade indicada acima.",
    "",
    "A primeira fatura vence 25 dias após a instalação;",
    "nos demais meses, o vencimento segue o ciclo de 30 em 30 dias.",
    "",
    "━━━━━━━━━━━━━━━━━━━━━",
    "✅ Confirma os dados do pedido?",
    "",
    "Toque em um dos botões abaixo:",
    "",
    "Parceiro oficial da Nio Fibra.",
    "SAC: 0800 001 1000 | WhatsApp: 21 3605-1000",
]

T2 = [
    "Olá, {{1}}!",
    "",
    "{{2}}, tudo bem?",
    "",
    "Parceiro oficial da Nio Fibra.",
    "Sua instalação da Nio Fibra está agendada para *{{3}}*,",
    "no período das *{{4}}*.",
    "",
    "Se você não puder estar presente, é necessário que",
    "uma pessoa maior de 18 anos esteja no local.",
    "",
    "A instalação é gratuita.",
    "Não realizamos instalações em dias de chuva.",
    "",
    "Toque em um dos botões abaixo:",
    "",
    "SAC: 0800 001 1000 | WhatsApp: 21 3605-1000",
]

T3 = [
    "Confirmação registrada. ✅",
    "",
    "Sua instalação Nio Fibra está confirmada para *{{1}}*,",
    "das *{{2}}*.",
    "",
    "O técnico entrará em contato por ligação e WhatsApp",
    "quando estiver a caminho.",
    "",
    "Se precisar de algo, use os botões abaixo.",
    "",
    "Parceiro oficial da Nio Fibra.",
]

T4 = [
    "Olá, {{1}}!",
    "",
    "{{2}},",
    "",
    "Parceiro oficial da Nio Fibra.",
    "Este é um lembrete da sua fatura Nio Fibra.",
    "",
    "Referência: *{{3}}*",
    "Valor: *{{4}}*",
    "Vencimento: *{{5}}* (em 5 dias).",
    "",
    "Toque em um dos botões abaixo para continuar.",
    "Se o pagamento já foi feito, escolha *Já paguei*.",
    "",
    "SAC: 0800 001 1000 | WhatsApp: 21 3605-1000",
]

T5 = [
    "Olá, {{1}}!",
    "",
    "{{2}},",
    "",
    "Parceiro oficial da Nio Fibra.",
    "Identificamos que sua fatura Nio Fibra está em atraso.",
    "",
    "Referência: *{{3}}*",
    "Valor: *{{4}}*",
    "Vencimento: *{{5}}* (há 5 dias).",
    "",
    "A confirmação do pagamento pode levar até 5 dias úteis.",
    "",
    "Toque em um dos botões abaixo para continuar.",
    "",
    "SAC: 0800 001 1000 | WhatsApp: 21 3605-1000",
]

T6 = [
    "Olá, {{1}}!",
    "",
    "{{2}},",
    "",
    "Parceiro oficial da Nio Fibra.",
    "Sua fatura Nio Fibra permanece em aberto.",
    "",
    "Referência: *{{3}}*",
    "Valor: *{{4}}*",
    "Vencimento: *{{5}}*",
    "Dias em atraso: *{{6}}*",
    "",
    "Toque em um dos botões abaixo para regularizar",
    "ou informar que o pagamento já foi feito.",
    "",
    "SAC: 0800 001 1000 | WhatsApp: 21 3605-1000",
]

T7 = [
    "Olá, {{1}}!",
    "",
    "{{2}},",
    "",
    "Parceiro oficial da Nio Fibra.",
    "Identificamos uma pendência no agendamento",
    "da sua instalação Nio Fibra.",
    "Na maioria dos casos isso não depende de você.",
    "",
    "Para reagendar, entre em contato pelo",
    "WhatsApp oficial da Nio: {{3}}.",
    "",
    "Toque em um dos botões abaixo se precisar",
    "de ajuda com o reagendamento ou tiver dúvidas.",
    "",
    "Obrigado por escolher a Nio Fibra.",
    "SAC: 0800 001 1000",
]

T8 = [
    "Olá, {{1}}!",
    "",
    "{{2}},",
    "",
    "Parceiro oficial da Nio Fibra.",
    "Sua instalação Nio Fibra foi concluída.",
    "Seja bem-vindo(a)!",
    "",
    "Guarde este WhatsApp para suporte e",
    "dúvidas sobre sua conexão.",
    "",
    "SAC: 0800 001 1000 | WhatsApp: 21 3605-1000",
    "",
    "Toque em um dos botões abaixo se precisar",
    "de ajuda agora.",
]


def main() -> None:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Pt(72)
        s.bottom_margin = Pt(72)
        s.left_margin = Pt(72)
        s.right_margin = Pt(72)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(t.add_run("Templates WhatsApp para aprovação Meta"), size=18, bold=True)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(
        st.add_run(
            "UTILITY · Botões Quick Reply · Confirmação · Instalação · Pendência · Boas-vindas · Cobrança"
        ),
        size=11,
        color=RGBColor(0x5B, 0x6B, 0x7D),
    )

    _h(doc, "Botões — regra Meta (obrigatório nestes templates)", 1)
    for x in [
        "Tipo: Quick Reply (resposta rápida). Não misturar com botão URL/telefone no mesmo template.",
        "Máximo: 3 botões por template · texto do botão ≤ 25 caracteres.",
        "No Manager: ao criar o template, adicione a seção Buttons → Quick reply.",
        "No corpo: orientar “Toque em um dos botões abaixo” — evitar “digite CORRETO/SIM”.",
        "O clique do cliente envia o texto do botão como mensagem; o CRM trata como CORRETO, "
        "CORRIGIR, Confirmar, etc.",
        "Opcional extra (fora destes 6): template só com Call phone 08000011000 — útil como SAC.",
    ]:
        _p(doc, f"• {x}")

    _h(doc, "Abertura padrão (todos os templates)", 1)
    _p(doc, "Em todos os corpos:")
    _code(
        doc,
        [
            "Olá, {{1}}!",
            "",
            "{{2}},",
        ],
    )
    doc.add_paragraph()
    _p(doc, "• {{1}} = saudação pelo horário de envio (CRM, fuso America/Sao_Paulo):", bold=True)
    _p(doc, "  05:00–11:59 → Bom dia | 12:00–17:59 → Boa tarde | 18:00–04:59 → Boa noite")
    _p(doc, "• {{2}} = primeiro nome do cliente (CRM) — sem “Sr(a).”")
    _p(doc, "Exemplo: Olá, Boa tarde! / Maria,")

    _h(doc, "Regras UTILITY", 1)
    for x in [
        "Categoria UTILITY (não Marketing).",
        "Sem oferta, desconto, “você ganha”, ameaça ou tom de cobrança abusiva.",
        "Cobrança = lembrete/atualização de fatura já existente (transacional).",
        "PIX/código de barras: enviar após o clique em “Quero a 2ª via” (janela 24h) ou documento.",
        "CPF/e-mail mascarados quando aparecerem (confirmação de pedido).",
        "Parceiro oficial da Nio Fibra · SAC 0800 001 1000 · WhatsApp 21 3605-1000.",
    ]:
        _p(doc, f"• {x}")

    _template(
        doc,
        titulo="1. Confirmação de pedido",
        nome="nio_confirmacao_pedido_v1",
        finalidade="Resumo do pedido + 3 botões de decisão",
        body=T1,
        botoes=["CORRETO", "CORRIGIR", "Falar com atendente"],
        variaveis=[
            "{{1}} saudação",
            "{{2}} primeiro nome",
            "{{3}} nome completo (campo Cliente)",
            "{{4}} CPF mascarado",
            "{{5}} e-mail mascarado",
            "{{6}} CEP",
            "{{7}} logradouro",
            "{{8}} número",
            "{{9}} complemento (ou \"-\")",
            "{{10}} bairro",
            "{{11}} cidade - UF",
            "{{12}} pagamento",
            "{{13}} plano + valor/mês",
            "{{14}} fidelidade",
        ],
        notas=[
            "CRM: CORRETO → segue fluxo; CORRIGIR → abre atendimento/edição; "
            "Falar com atendente → fila humana.",
            "{{2}} = primeiro nome; {{3}} = nome completo no bloco Cliente.",
            "Se a Meta limitar variáveis, unir endereço em um único {{n}}.",
        ],
    )

    _template(
        doc,
        titulo="2. Lembrete de instalação",
        nome="nio_lembrete_instalacao_v1",
        finalidade="Aviso da agenda + confirmar / reagendar / suporte",
        body=T2,
        botoes=["Confirmar", "Reagendar", "Suporte"],
        variaveis=[
            "{{1}} saudação",
            "{{2}} primeiro nome",
            "{{3}} data",
            "{{4}} período/horário",
        ],
        notas=[
            "CRM: Confirmar → envia template 3; Reagendar → pede nova data na janela 24h; "
            "Suporte → fila humana.",
            "Aceitar também texto livre SIM/SUPORTE como fallback se o cliente digitar.",
        ],
    )

    _template(
        doc,
        titulo="3. Instalação confirmada (após Confirmar)",
        nome="nio_instalacao_confirmada_v1_2",
        finalidade="Confirmar agenda (continuação do lembrete — sem repetir saudação/nome)",
        body=T3,
        botoes=["Entendi", "Reagendar", "Suporte"],
        variaveis=[
            "{{1}} data",
            "{{2}} horário",
        ],
        notas=[
            "Sem Olá/saudação: é resposta imediata após Confirmar no template 2.",
            "CRM: Entendi → encerra/ok; Reagendar → abre reagendamento; Suporte → fila humana.",
        ],
    )

    _h(doc, "Cobrança (fatura)", 1)
    _p(
        doc,
        "Três templates distintos. Em todos: 3 botões Quick Reply para 2ª via, "
        "baixa de pagamento ou suporte — sem pedir para digitar.",
    )

    _botoes_cobranca = ["Quero a 2ª via", "Já paguei", "Falar com suporte"]
    _notas_btn_cob = [
        "CRM: Quero a 2ª via → envia PIX/barras/PDF na janela 24h; "
        "Já paguei → pede comprovante ou marca análise; Falar com suporte → fila humana.",
    ]

    _template(
        doc,
        titulo="4. Lembrete — 5 dias antes do vencimento",
        nome="nio_fatura_lembrete_5d_antes_v1",
        finalidade="Avisar fatura a vencer em 5 dias + 3 botões",
        body=T4,
        botoes=_botoes_cobranca,
        variaveis=[
            "{{1}} saudação",
            "{{2}} primeiro nome",
            "{{3}} referência da fatura",
            "{{4}} valor",
            "{{5}} data de vencimento",
        ],
        notas=[
            "Disparo CRM: data_vencimento - 5 dias.",
            *_notas_btn_cob,
        ],
    )

    _template(
        doc,
        titulo="5. Aviso — 5 dias após o vencimento",
        nome="nio_fatura_vencida_5d_v1",
        finalidade="Avisar fatura vencida há 5 dias + 3 botões",
        body=T5,
        botoes=_botoes_cobranca,
        variaveis=[
            "{{1}} saudação",
            "{{2}} primeiro nome",
            "{{3}} referência da fatura",
            "{{4}} valor",
            "{{5}} data de vencimento",
        ],
        notas=[
            "Disparo CRM: data_vencimento + 5 dias, se ainda unpaid.",
            *_notas_btn_cob,
        ],
    )

    _template(
        doc,
        titulo="6. Cobrança recorrente — fatura ainda em aberto",
        nome="nio_fatura_cobranca_recorrente_v1",
        finalidade="Lembretes periódicos + 3 botões enquanto unpaid",
        body=T6,
        botoes=_botoes_cobranca,
        variaveis=[
            "{{1}} saudação",
            "{{2}} primeiro nome",
            "{{3}} referência da fatura",
            "{{4}} valor",
            "{{5}} data de vencimento",
            "{{6}} dias em atraso (número)",
        ],
        notas=[
            "Disparo CRM sugerido: a cada 7 dias após o template 5, enquanto unpaid "
            "(ex.: D+12, D+19, D+26…), com teto de tentativas definido pela operação.",
            "Não usar linguagem de ameaça, negativação ou oferta — mantém UTILITY.",
            "Parar automaticamente ao baixar o pagamento.",
            *_notas_btn_cob,
        ],
    )

    _h(doc, "Agenda sugerida no CRM (cobrança)", 1)
    _p(doc, "• D−5: template 4 (lembrete antes)")
    _p(doc, "• D+5: template 5 (vencida)")
    _p(doc, "• D+12, D+19, D+26…: template 6 (recorrente), até pagar ou atingir limite")
    _p(doc, "• Clique em “Quero a 2ª via” (ou qualquer resposta): janela 24h → enviar PIX/barras/PDF")

    _h(doc, "Pendência e boas-vindas (esteira)", 1)

    _template(
        doc,
        titulo="7. Pendência de agendamento (tipo Cliente)",
        nome="nio_pendencia_reagendamento_v1",
        finalidade="Avisar pendência + orientar reagendamento no canal oficial Nio",
        body=T7,
        botoes=["Reagendar", "Falar com atendente", "Entendi"],
        variaveis=[
            "{{1}} saudação",
            "{{2}} primeiro nome",
            "{{3}} WhatsApp oficial Nio (ex.: 21 3605-1000)",
        ],
        notas=[
            "Disparo CRM: esteira, ao registrar pendência tipo CLIENTE, "
            "somente após modal «Enviar WhatsApp ao cliente?» = Sim.",
            "CRM: Reagendar → coleta intenção / orienta canal Nio; "
            "Falar com atendente → fila humana; Entendi → encerra ok.",
            "Não culpar o cliente; tom UTILITY (atualização de status do pedido).",
        ],
    )

    _template(
        doc,
        titulo="8. Boas-vindas pós-instalação",
        nome="nio_boas_vindas_v1",
        finalidade="Boas-vindas ao concluir instalação (Esteira / status INSTALADA)",
        body=T8,
        botoes=["Entendi", "Falar com atendente", "Suporte"],
        variaveis=[
            "{{1}} saudação",
            "{{2}} primeiro nome",
        ],
        notas=[
            "Disparo CRM: ao virar INSTALADA (automático), com opção de não enviar.",
            "CRM: Entendi → ok; Falar com atendente / Suporte → fila humana.",
        ],
    )

    _h(doc, "Mapa rápido — texto do botão → ação CRM", 1)
    for row in [
        "CORRETO → avança pedido",
        "CORRIGIR / Reagendar → coleta correção ou nova data (janela 24h)",
        "Confirmar → envia nio_instalacao_confirmada_v1_2",
        "Entendi → encerra interação ok",
        "Quero a 2ª via → envia fatura (PIX/barras/PDF)",
        "Já paguei → solicita comprovante / fila de baixa",
        "Suporte / Falar com suporte / Falar com atendente → fila humana",
    ]:
        _p(doc, f"• {row}")

    _h(doc, "Checklist de submissão", 1)
    for item in [
        "Categoria = UTILITY em todos",
        "Buttons = Quick reply (3 botões) em todos os templates",
        "Incluir templates 7 (pendência) e 8 (boas-vindas) junto dos 6 já aprovados",
        "Idioma pt_BR · sem marketing/oferta",
        "Após aprovação na WABA (#194), validar envio no CRM (esteira + scheduler)",
        "Corpo pede “Toque no(s) botão(ões)” — sem “digite X”",
        "Todos com Olá, {{1}}! + {{2}} primeiro nome",
        "Sem Sr(a). / sem tom promocional ou ameaçador",
        "Amostra sem dados reais",
        "Após APPROVED: Número B (#194)",
    ]:
        _p(doc, f"☐ {item}")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Gerado: {OUT}")


if __name__ == "__main__":
    main()
